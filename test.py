from docx import Document

# Create a new Document
doc = Document()

# Add cover page
doc.add_paragraph('Library Management System Report', style='Title')
doc.add_paragraph('By: Your Name')

# Add contents page
doc.add_page_break()
doc.add_paragraph('Contents', style='Heading 1')
doc.add_paragraph('1. Introduction')
doc.add_paragraph('2. Hardware Costs')
doc.add_paragraph('3. Software Costs')
doc.add_paragraph('4. Summary of Estimated Costs')
doc.add_paragraph('5. References')

# Add Introduction
doc.add_page_break()
doc.add_paragraph('1. Introduction', style='Heading 1')
doc.add_paragraph('''
This report provides an analysis of the capital costs involved in setting up a computer system 
for a library management system. The system is designed to replace a manual card-based cataloging 
and tracking system with a digital solution. The software for the library management system will be 
custom-built using Python with a GUI, and will utilize free SQL software for the database management system.

The primary goal is to keep costs as low as possible while ensuring the system meets the library's needs.
''')

# Add Hardware Costs
doc.add_page_break()
doc.add_paragraph('2. Hardware Costs', style='Heading 1')
doc.add_paragraph('''
The hardware required for the library management system consists of a computer (desktop or laptop), 
peripherals like a printer, and optional backup storage devices. Here are the estimated costs:

- Budget Option (Desktop): £300 - £500
- Mid-range Option (Desktop/Laptop): £500 - £700
- High-end Option (Laptop/Desktop with higher performance): £700 - £1,000

Optional Peripherals:
- External Backup Drive: £50 - £100
''')

# Add Software Costs
doc.add_page_break()
doc.add_paragraph('3. Software Costs', style='Heading 1')
doc.add_paragraph('''
The software used for the library management system will be a Python program with a GUI, 
running on free SQL software for database management. The following free software stack will be used:

- Python (free and open-source): £0
- MySQL or SQLite (free and open-source): £0
- GUI Framework (Tkinter, PyQt, or Kivy): £0
- Development Tools (e.g., VS Code, PyCharm Community Edition): £0
- Operating System (Windows or Linux): £0 - £150 (if opting for a Windows license)

Optional:
- Paid Antivirus/Backup Software: £30 - £60 per year
''')

# Add Summary of Estimated Costs
doc.add_page_break()
doc.add_paragraph('4. Summary of Estimated Costs', style='Heading 1')
doc.add_paragraph('''
Here is a summary of the estimated initial capital costs:

- **Budget Setup (using free software stack):**
  - **Hardware:** £300 - £500
  - **Software:** £0 (using open-source tools)
  - **Total Budget Cost:** £300 - £500

- **Mid-Range Setup (with optional paid software):**
  - **Hardware:** £500 - £900
  - **Software:** £130 - £210 (if opting for Windows and antivirus)
  - **Total Mid-Range Cost:** £630 - £1,110

- **High-End Setup (high performance and optional paid software):**
  - **Hardware:** £900 - £1,200
  - **Software:** £130 - £210 (if opting for Windows and antivirus)
  - **Total High-End Cost:** £1,030 - £1,410
''')

# Add References
doc.add_page_break()
doc.add_paragraph('5. References', style='Heading 1')
doc.add_paragraph('''
1. Library Management System Overview. (n.d.). Retrieved from https://example.com
2. Python Documentation. (n.d.). Retrieved from https://python.org
3. MySQL Documentation. (n.d.). Retrieved from https://mysql.com
4. Open Source GUI Frameworks for Python. (n.d.). Retrieved from https://example.com
''')

# Save the Word document
doc_file_path = 'Desktop'
doc.save(doc_file_path)

# Provide the file path for download
doc_file_path
